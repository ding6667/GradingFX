import json
import os
import sys  # 💡 必须导入 sys 模块，否则无法接收 sys.argv 命令行参数
import uuid
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


# ==================== 样式美化辅助函数 ====================
def set_cell_background(cell, color_hex):
    """设置单元格背景颜色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """设置单元格内边距（防止表格太挤）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def parse_markdown_to_docx(doc, md_text):
    """核心算法：解析带有Markdown表格的文本并写入Word"""
    lines = md_text.strip().split('\n')
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # 识别 Markdown 表格行
        if stripped.startswith('|') and stripped.endswith('|'):
            # 过滤掉表格的分割线如 | :--- | :---: |
            if '---' in stripped:
                continue
            in_table = True
            # 提取单元格数据
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            # 如果之前在表格中，现在退出了，先把表格画出来
            if in_table and table_data:
                draw_word_table(doc, table_data)
                table_data = []
                in_table = False

            # 处理普通的文本、标题
            if stripped.startswith('# '):
                p = doc.add_paragraph()
                r = p.add_run(stripped.replace('# ', ''))
                r.font.size = Pt(16)
                r.bold = True
                r.font.color.rgb = RGBColor(31, 78, 121)  # 深蓝色大标题
                p.paragraph_format.space_before = Pt(12)
            elif stripped.startswith('## '):
                p = doc.add_paragraph()
                r = p.add_run(stripped.replace('## ', ''))
                r.font.size = Pt(13)
                r.bold = True
                r.font.color.rgb = RGBColor(46, 116, 181)  # 浅蓝色二级标题
                p.paragraph_format.space_before = Pt(8)
            elif stripped.startswith('### '):
                p = doc.add_paragraph()
                r = p.add_run(stripped.replace('### ', ''))
                r.font.size = Pt(11)
                r.bold = True
                p.paragraph_format.space_before = Pt(6)
            elif stripped.startswith('**') and stripped.endswith('**'):
                p = doc.add_paragraph()
                r = p.add_run(stripped.replace('**', ''))
                r.bold = True
            elif stripped == "---":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("—" * 30).font.color.rgb = RGBColor(200, 200, 200)
            elif stripped:
                # 简单替换粗体语法 **文本** 为 Word 粗体
                p = doc.add_paragraph()
                parts = stripped.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:  # 奇数部分是包裹在 ** 内部的
                        run.bold = True

    # 循环结束后，如果最后是表格，补画
    if table_data:
        draw_word_table(doc, table_data)


def draw_word_table(doc, table_data):
    """将提取出来的二维数组绘制成美化的 Word 原生表格（具备超强容错，防止列数不匹配崩溃）"""
    rows = len(table_data)
    if rows == 0: return

    # 找出这组表格数据中，所有行里“最大的列数”
    cols = max(len(row) for row in table_data)
    if cols == 0: return

    # 创建对应行和最大列数的 Word 表格
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = True
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(table_data):
        # 如果当前行的列数比最大列数少，用空字符串补齐
        while len(row_data) < cols:
            row_data.append("")

        row = table.rows[r_idx]
        for c_idx, cell_value in enumerate(row_data):
            if c_idx >= len(row.cells):
                break

            cell = row.cells[c_idx]
            cell.text = cell_value
            set_cell_margins(cell)  # 加宽间距

            # 美化表头（第一行）
            if r_idx == 0:
                set_cell_background(cell, "2E74B5")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # 白字
            else:
                # 斑马纹：偶数行设置极淡的灰色背景
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F2F2F2")

    doc.add_paragraph()  # 表格后留空行


# ==================== 接收临时文件路径的核心处理接口 ====================
def process_json_file(file_path, output_dir):
    """
    接收来自 Java 的临时 JSON 文件路径，读取并将其转换为规整的 Word 文档
    """
    # 1. 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"错误：找不到临时的 JSON 文件 -> {file_path}")
        return False

    try:
        # 2. 从临时文件中以 UTF-8 编码安全读取 JSON 数据
        with open(file_path, "r", encoding="utf-8") as f:
            data_dict = json.load(f)
    except Exception as e:
        print(f"JSON 文件读取或解析失败: {str(e)}")
        return False

    # 3. 检查数据中是否存在需要的字段
    if "output" not in data_dict:
        print("JSON 数据中未找到 'output' 键，无法提取批阅结果。")
        return False

    raw_outputs = data_dict["output"]

    # 兼容处理：如果是字符串，则转为包含单个元素的列表；如果是列表，直接使用
    if isinstance(raw_outputs, str):
        reports = [raw_outputs]
    elif isinstance(raw_outputs, list):
        reports = raw_outputs
    else:
        print("'output' 字段类型不支持，需要 str 或 list")
        return False

    # 4. 创建 Word 对象并配置基础样式
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    print(f"开始解析大模型批阅报告，共找到 {len(reports)} 篇...")
    for idx, report_content in enumerate(reports, 1):
        print(f"正在转换第 {idx} 篇报告...")
        # 解析 Markdown 格式并将段落和表格写入 Word
        parse_markdown_to_docx(doc, report_content)
        # 每篇报告之间加一个换页符，让一章占一页
        if idx < len(reports):
            doc.add_page_break()

    # 5. 自动创建并保存到你指定的系统路径下
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)  # 如果文件夹不存在，自动创建

    output_path = os.path.join(output_dir, f"学生Java作业批阅汇总报告_{uuid.uuid4().hex}.docx")
    doc.save(output_path)
    print(f"处理成功！Word 表格报告已生成至-{output_path}")
    return True

if __name__ == "__main__":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    # 确保命令行参数存在
    if len(sys.argv) < 3:
        print("错误：未检测到传入的临时文件路径或输出路径参数！")
        sys.exit(1)

    # sys.argv[1] 接收的是 Java 传过来的临时文件路径
    temp_file_path = sys.argv[1]
    output_dir = sys.argv[2]
    process_json_file(temp_file_path, output_dir)
